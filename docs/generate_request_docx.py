# -*- coding: utf-8 -*-
"""서풍 홈페이지 — 고객사 정보 요청서 (Word 문서 생성)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUTPUT = Path(__file__).parent / "INFO_REQUEST.docx"

# ── 색상
COLOR_OCEAN = RGBColor(0x0E, 0x74, 0x90)
COLOR_GOLD = RGBColor(0xE8, 0xA3, 0x17)
COLOR_NAVY = RGBColor(0x0A, 0x16, 0x28)
COLOR_GRAY = RGBColor(0x6B, 0x76, 0x88)
COLOR_LIGHT_GRAY = RGBColor(0xE5, 0xE8, 0xEC)

KOR_FONT = "맑은 고딕"


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell, color_hex="cfd6df", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color_hex)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def style_run(run, *, size=10.5, bold=False, color=None, font=KOR_FONT):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=18, bold=True, color=COLOR_NAVY)
    # 하단 라인
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "0e7490")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=13.5, bold=True, color=COLOR_OCEAN)
    return p


def add_para(doc, text, *, size=10.5, bold=False, italic=False, color=None, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, color=color)
    run.italic = italic
    return p


def add_info_box(doc, text):
    """안내 박스 (회색 배경)"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, "f3f5f7")
    set_cell_border(cell, color_hex="cfd6df", size="6")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    style_run(run, size=10, color=COLOR_GRAY)
    return table


def add_table_with_header(doc, headers, rows, col_widths_cm=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = False

    # 헤더
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], "0e7490")
        set_cell_border(hdr[i])
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        style_run(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 데이터
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            set_cell_border(cells[i])
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val) if val is not None else "")
            style_run(run, size=10)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 컬럼 너비
    if col_widths_cm:
        for col_idx, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[col_idx].width = Cm(width)

    # 표 뒤 빈 줄
    doc.add_paragraph()
    return table


def add_checkbox_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("☐  ")
    style_run(run, size=11, color=COLOR_OCEAN)
    run2 = p.add_run(text)
    style_run(run2, size=10.5)


# ──────────────────────────────────────────────
def build():
    doc = Document()

    # 기본 스타일
    style = doc.styles["Normal"]
    style.font.name = KOR_FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), KOR_FONT)

    # 페이지 여백
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── 표지 영역
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("서풍 홈페이지")
    style_run(run, size=12, color=COLOR_OCEAN, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(20)
    run = subtitle.add_run("고객사 정보 요청서")
    style_run(run, size=24, bold=True, color=COLOR_NAVY)
    # 하단 굵은 라인
    pPr = subtitle._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "24")
    bottom.set(qn("w:color"), "e8a317")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # 메타 정보
    meta = doc.add_table(rows=4, cols=2)
    meta_data = [
        ("작성일", "2026-05-11"),
        ("대상", "영어조합법인 서풍"),
        ("회신", "gpt1@dure-coop.or.kr"),
        ("용도", "seopung.co.kr 기업 홈페이지 최종 납품 전 정보 확정"),
    ]
    for i, (label, value) in enumerate(meta_data):
        c1 = meta.cell(i, 0)
        c2 = meta.cell(i, 1)
        c1.width = Cm(2.5)
        c2.width = Cm(14)
        c1.text = ""
        c2.text = ""
        p1 = c1.paragraphs[0]
        run1 = p1.add_run(label)
        style_run(run1, size=10, bold=True, color=COLOR_OCEAN)
        p2 = c2.paragraphs[0]
        run2 = p2.add_run(value)
        style_run(run2, size=10)

    doc.add_paragraph()

    # ── 안내
    add_heading_2(doc, "안내")
    add_info_box(
        doc,
        "홈페이지 골격과 디자인은 완성되어 있습니다. 아래 정보가 확정되어 회신되면, "
        "단일 설정 파일과 사진 파일만 교체하면 사이트 전체에 자동 반영됩니다. "
        "코드 수정 없이 정보·사진 입력만으로 완성됩니다.",
    )
    add_para(doc, "작성 원칙", bold=True, size=11)
    add_checkbox_line(doc, "확정된 항목만 작성해주세요. 미확정은 비워두면 사이트에서 자동으로 숨겨집니다.")
    add_checkbox_line(doc, "사진은 가능한 한 고품질로 부탁드립니다 (가로 1,600px 이상 권장).")
    add_checkbox_line(doc, "인증서·사업자등록증 등 문서는 PDF 또는 컬러 스캔본으로 첨부.")

    doc.add_paragraph()

    # ── A. 법인 기본 정보
    add_heading_1(doc, "A. 법인 기본 정보 (필수)")
    add_table_with_header(
        doc,
        ["#", "항목", "현재 사이트 표기", "회신란"],
        [
            ["A-1", "정식 법인명", "영어조합법인 서풍", "□ 그대로  □ 정정: ____________________"],
            ["A-2", "대표자 성명", "김태환", "□ 그대로  □ 정정: ____________________"],
            ["A-3", "대표자 직함", "(미표기)", "(예: 대표이사 / 상무이사) ____________________"],
            ["A-4", "법인 설립일", "(미표기)", "YYYY-MM-DD ____________________"],
            ["A-5", "사업자등록번호", "417-81-41979", "□ 그대로  □ 정정: ____________________"],
            ["A-6", "본사 주소 (도로명)", "전라남도 여수시 석교로 121", "□ 그대로  □ 정정: ____________________"],
            ["A-7", "상세 주소", "화양면", "____________________"],
            ["A-8", "우편번호", "(미표기)", "____________________"],
            ["A-9", "대표 전화번호", "061-686-0508", "□ 그대로  □ 정정: ____________________"],
            ["A-10", "FAX 번호", "(미표기)", "____________________"],
            ["A-11", "공식 이메일", "seopung@naver.com", "□ 그대로  □ 정정 (도메인 이메일 권장): ____________________"],
        ],
        col_widths_cm=[1.2, 3.5, 4.5, 7.5],
    )

    add_para(doc, "첨부 요청", bold=True, size=11)
    add_checkbox_line(doc, "사업자등록증 PDF")
    add_checkbox_line(doc, "법인 인감증명서 (필요시)")
    doc.add_paragraph()

    # ── B. 매출·실적
    add_heading_1(doc, "B. 회사 매출·실적 (선택, 가능 시 응답)")
    add_info_box(
        doc,
        "모든 숫자는 사이트의 '숫자로 보는 서풍' 섹션 등에 노출됩니다. "
        "확정된 수치만 응답해 주세요. 미확정 수치는 사이트에 표시되지 않습니다.",
    )

    add_table_with_header(
        doc,
        ["#", "항목", "회신란"],
        [
            ["B-1", "2025년 실제 매출액", "(예: 379억원) ____________________"],
            ["B-2", "2026년 목표 매출액", "(예: 400억원) ____________________"],
            ["B-3", "누적 OEM 개발 품목 수", "(예: 134개) ____________________"],
            ["B-4", "현재 양산 운영 품목 수", "(예: 66개) ____________________"],
            ["B-5", "2026년 설비투자 규모", "(예: 9억원+) ____________________"],
            ["B-6", "평균 신제품 개발기간", "(예: 5개월) ____________________"],
            ["B-7", "전체 직원 수", "____________________"],
            ["B-8", "정기 거래처 수", "____________________"],
        ],
        col_widths_cm=[1.2, 5.5, 10.0],
    )

    add_para(doc, "품질 KPI (인증·품질 페이지 노출)", bold=True, size=11)
    add_table_with_header(
        doc,
        ["#", "항목", "회신란"],
        [
            ["B-9", "불량률", "(예: 0.02%) ____________________"],
            ["B-10", "납기 준수율", "(예: 99.7%) ____________________"],
            ["B-11", "거래처 재계약률", "(예: 95%+) ____________________"],
            ["B-12", "일일 처리량", "(예: 15톤) ____________________"],
        ],
        col_widths_cm=[1.2, 5.5, 10.0],
    )

    # ── C. 인증 정보
    add_heading_1(doc, "C. 인증 정보 (핵심)")
    add_info_box(
        doc,
        "사이트의 인증 페이지는 회사의 최대 강점으로 강조되어 있습니다. "
        "정확한 인증번호와 발급일을 채워주시면 '인증서 다운로드' 섹션이 활성화됩니다.",
    )
    add_table_with_header(
        doc,
        ["#", "인증명", "인증번호", "최초 발급일", "PDF 첨부"],
        [
            ["C-1", "HACCP", "____________", "YYYY-MM-DD", "☐"],
            ["C-2", "수산물 이력추적관리", "____________", "YYYY-MM-DD", "☐"],
            ["C-3", "수산물 품질인증", "____________", "YYYY-MM-DD", "☐ / 미보유 ☐"],
            ["C-4", "ASC (양식 지속가능성)", "____________", "YYYY-MM-DD", "☐"],
            ["C-5", "MSC (어업 지속가능성)", "____________", "YYYY-MM-DD", "☐"],
            ["C-6", "FSSC 22000 (추진 예정)", "예정", "예정 일자: ______", "(해당없음)"],
            ["C-7", "기타 보유 인증", "____________", "YYYY-MM-DD", "☐"],
        ],
        col_widths_cm=[1.2, 4.5, 3.5, 3.0, 2.5],
    )

    # ── D. 연혁
    add_heading_1(doc, "D. 회사 연혁")
    add_para(doc, "사이트의 연혁 타임라인에 표시됩니다. 확정된 항목만 작성해주세요.", size=10, color=COLOR_GRAY)

    add_table_with_header(
        doc,
        ["연도", "주요 마일스톤"],
        [["____", "______________________________________________________"]] * 8,
        col_widths_cm=[2.0, 14.0],
    )

    add_para(doc, "참고로 사이트에 임시 기재되어 있던 항목 (사실 확인 필요)", bold=True, size=10.5)
    legacy_history = [
        "2008 — HACCP 최초 인증",
        "2011 — HACCP 기반 품질위생관리 체계 안정화",
        "2013 — 수산물 이력추적관리 시스템 도입",
        "2017 — '레몬담은수산물' 브랜드 출시",
        "2019 — '마리네이드수산물' 브랜드 출시",
        "2024 — ASC·MSC 인증, 8억원 설비 투자",
        "2025 — 7.4억원 설비 투자",
        "2026 — AI 엑스레이·초분광 도입, FSSC 22000 추진",
    ]
    for item in legacy_history:
        add_checkbox_line(doc, item + "   [□ 사용 / □ 수정 / □ 삭제]")

    doc.add_paragraph()

    # ── E. 취급 어종
    add_heading_1(doc, "E. 취급 어종")
    add_para(doc, "사이트 contact 페이지의 '취급어종' 항목에 표시됩니다.", size=10, color=COLOR_GRAY)
    add_table_with_header(
        doc,
        ["#", "회신란"],
        [["E-1", "정식 취급 어종 리스트 (쉼표로 구분): ____________________"]],
        col_widths_cm=[1.2, 14.8],
    )
    add_para(doc, "참고로 사이트에 임시 기재되어 있던 어종", bold=True, size=10.5)
    add_para(doc, "참조기 / 삼치 / 오징어 / 갈치 / 고등어 / 아귀 / 방어 / 달고기 / 붕장어", size=10, color=COLOR_GRAY)

    doc.add_paragraph()

    # ── F. 거래처
    add_heading_1(doc, "F. 거래처·파트너사 (로고 사용 동의 필요)")
    add_info_box(
        doc,
        "사이트의 'Trusted Partners' 섹션에 표시. 각 거래처의 로고 사용 서면 동의가 필요합니다. "
        "동의 받지 못한 곳은 표시되지 않습니다.",
    )
    add_table_with_header(
        doc,
        ["#", "거래처명", "사용 동의", "로고 파일 첨부"],
        [
            ["F-1", "____________________", "□ 받음  □ 미확보", "□ 첨부 (PNG/SVG)"],
            ["F-2", "____________________", "□ 받음  □ 미확보", "□ 첨부"],
            ["F-3", "____________________", "□ 받음  □ 미확보", "□ 첨부"],
            ["F-4", "____________________", "□ 받음  □ 미확보", "□ 첨부"],
            ["F-5", "____________________", "□ 받음  □ 미확보", "□ 첨부"],
        ],
        col_widths_cm=[1.2, 6.0, 4.5, 4.3],
    )
    add_para(
        doc,
        "참고: 이전 사이트에 텍스트로 들어가 있던 풀무원 / 푸드머스 / 홈플러스 / 이마트 / 쿠팡 등은 "
        "모두 로고 사용 동의가 확인되지 않아 임시로 placeholder 처리되어 있습니다.",
        size=10,
        color=COLOR_GRAY,
        italic=True,
    )

    doc.add_paragraph()

    # ── G. 관계사
    add_heading_1(doc, "G. 관계사 / 그룹사")
    add_para(doc, "사이트 contact 페이지 '관계사 정보' 섹션에 표시. 사용 동의 받은 법인만 등록.", size=10, color=COLOR_GRAY)
    add_table_with_header(
        doc,
        ["#", "법인명", "역할", "한 줄 설명", "사용 동의"],
        [
            ["G-1", "____________", "Manufacturing", "____________", "□"],
            ["G-2", "____________", "Distribution", "____________", "□"],
            ["G-3", "____________", "Storage", "____________", "□"],
            ["G-4", "____________", "Purchase", "____________", "□"],
        ],
        col_widths_cm=[1.2, 4.0, 3.0, 6.0, 1.8],
    )
    add_para(doc, "참고로 사이트에 임시 기재되어 있던 관계사 (사실 확인 필요)", bold=True, size=10.5)
    for item in [
        "영어조합법인 서풍 (Manufacturing)",
        "㈜여수유통 (Distribution)",
        "㈜대주냉장 (Storage)",
        "중매인 49호 (Purchase)",
    ]:
        add_checkbox_line(doc, item + "   [□ 사용 / □ 수정 / □ 삭제]")

    doc.add_paragraph()

    # ── H. 사진
    add_heading_1(doc, "H. 사진 자료 ★")
    add_info_box(
        doc,
        "현재 사이트의 (사진필요) 표시된 자리는 모두 placeholder입니다. "
        "아래 사진들이 회신되면 placeholder가 자동 교체됩니다.",
    )
    add_para(doc, "필수 사진 (H1~H5)", bold=True, size=11)
    add_table_with_header(
        doc,
        ["#", "위치", "요구 사양", "제출"],
        [
            ["H-1", "대표 인사말 프로필", "정장 또는 깔끔한 작업복, 세로 4:5, 1,600×2,000px 이상", "☐"],
            ["H-2", "공장 외관 / 대표 컷", "스튜디오 품질, 가로 16:9 또는 4:5, 깔끔한 배경", "☐"],
            ["H-3", "생산공정 6단계", "원료수매/입고검사/해동·전처리/가공/동결/포장 — 단계별 1장", "☐"],
            ["H-4", "AI 엑스레이 / 핵심 설비", "2026년 도입 예정 설비 카탈로그 컷이라도 가능", "☐"],
            ["H-5", "인증서 액자 / 현판", "HACCP·ASC·MSC 등 사무실 벽에 걸린 인증서 사진", "☐"],
        ],
        col_widths_cm=[1.2, 4.5, 8.5, 1.8],
    )

    add_para(doc, "권장 사진 (H6~H10)", bold=True, size=11)
    add_table_with_header(
        doc,
        ["#", "위치", "요구 사양", "제출"],
        [
            ["H-6", "주요 OEM 제품 패키지", "4종 (냉동수산/굴비/밀키트/선물세트), 정사각형 또는 4:3", "☐"],
            ["H-7", "거래처 로고 PNG/SVG", "F 섹션 사용 동의 받은 곳만, 투명배경", "☐"],
            ["H-8", "회사 단체 사진", "정렬된 단정한 컷, 가로 16:9", "☐"],
            ["H-9", "공장 내부 라인", "위생복 착용 작업 모습, 가로 4:3", "☐"],
            ["H-10", "회사 로고 SVG (벡터)", "현재 PNG만 보유 — 벡터 원본", "☐"],
        ],
        col_widths_cm=[1.2, 4.5, 8.5, 1.8],
    )
    add_para(
        doc,
        "사진 파일은 첨부 또는 클라우드 링크로 회신 → 작업자가 정리해서 업로드합니다.",
        size=10,
        color=COLOR_GRAY,
        italic=True,
    )

    doc.add_paragraph()

    # ── I. 기타
    add_heading_1(doc, "I. 기타 운영 정보 (선택)")
    add_table_with_header(
        doc,
        ["#", "항목", "회신란"],
        [
            ["I-1", "카카오톡 채널 ID", "____________________"],
            ["I-2", "인스타그램 URL", "____________________"],
            ["I-3", "유튜브 채널 URL", "____________________"],
            ["I-4", "네이버 블로그 URL", "____________________"],
            ["I-5", "명함 QR 인쇄용 URL", "https://seopung.co.kr (그대로 사용 가능)"],
        ],
        col_widths_cm=[1.2, 5.0, 9.8],
    )

    # ── J. 문구 검토
    add_heading_1(doc, "J. 페이지 별 문구 검토 (선택)")
    add_para(doc, "디자이너가 임시 작성한 문구. 회사 톤과 맞지 않는 부분이 있으면 표시해주세요.", size=10, color=COLOR_GRAY)
    add_table_with_header(
        doc,
        ["#", "위치", "현재 문구", "의견"],
        [
            ["J-1", "홈 헤드라인", "수산 가공의 새로운 기준, 서풍이 만들어갑니다", "□ 좋음  □ 수정: ______"],
            ["J-2", "홈 서브타이틀", "HACCP · ASC · MSC 글로벌 인증 기반의 프리미엄 수산 OEM 파트너", "□ 좋음  □ 수정: ______"],
            ["J-3", "회사 슬로건", "바다의 가치를 세상의 식탁으로", "□ 좋음  □ 수정: ______"],
            ["J-4", "인증 페이지 헤드", "글로벌 기준을 넘어, 새로운 기준을 만듭니다", "□ 좋음  □ 수정: ______"],
            ["J-5", "OEM 약속 #1", "문의 접수 24시간 이내 담당자가 직접 회신드립니다", "□ 그대로  □ 수정: ______"],
            ["J-6", "OEM 약속 #3", "제품 검토를 위한 샘플을 무상 제공합니다", "□ 그대로  □ 수정: ______"],
        ],
        col_widths_cm=[1.0, 3.5, 7.0, 4.5],
    )

    # ── 회신 방법
    add_heading_1(doc, "회신 방법")
    add_checkbox_line(doc, "본 문서에 직접 답변을 적은 후 PDF/Word 등으로 회신")
    add_checkbox_line(doc, "또는 카카오톡으로 항목 번호만 적어 회신 (예: A-3: 대표이사, B-1: 379억)")
    add_checkbox_line(doc, "사진·문서 파일은 이메일 첨부 또는 클라우드 링크")
    doc.add_paragraph()

    add_info_box(
        doc,
        "회신 후 24시간 이내에 사이트에 반영되어 검수 가능한 미리보기 URL을 보내드립니다.\n"
        "문의: gpt1@dure-coop.or.kr",
    )

    # 푸터
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "본 문서는 홈페이지 납품 전 데이터 정확성 확보를 위한 정보 요청서입니다. "
        "응답하지 않은 항목은 사이트에서 자동으로 숨겨지며, 추후 확정 시 즉시 반영됩니다."
    )
    style_run(run, size=9, color=COLOR_GRAY, font=KOR_FONT)
    run.italic = True

    doc.save(OUTPUT)
    print(f"Wrote: {OUTPUT}")


if __name__ == "__main__":
    build()
