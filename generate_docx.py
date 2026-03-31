"""영어조합법인 서풍 홈페이지 납품설명서 Word 문서 생성"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '맑은 고딕'
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    h.font.color.rgb = RGBColor(0x1a, 0x2a, 0x4a)

# ── 페이지 여백 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        # Header bg
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1a2a4a',
            qn('w:val'): 'clear',
        })
        shading.append(bg)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('영어조합법인 서풍')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x2a, 0x4a)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('기업 홈페이지 납품 설명서')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0xc8, 0x96, 0x2d)

doc.add_paragraph()
doc.add_paragraph()

info_items = [
    ('문서번호', 'SP-WEB-2026-001'),
    ('작성일', '2026년 3월 24일'),
    ('버전', '1.0'),
    ('접속주소', 'https://durecoop.github.io/seopung-web/'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(value)
    run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('납품처: 영어조합법인 서풍')
run.font.size = Pt(12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('작성: 두레쿱')
run.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════
# 목차
# ══════════════════════════════════════════════
doc.add_heading('목 차', level=1)
toc_items = [
    '1. 프로젝트 개요',
    '2. 기술 사양',
    '3. 사이트 구조 (전체 12페이지)',
    '   3.3 페이지별 스크린샷',
    '   3.4 모바일 반응형 스크린샷',
    '4. 주요 기능',
    '5. 콘텐츠 현황',
    '6. 요청 자료 목록 (최종 완성에 필요)',
    '7. 기존 사진 퀄리티 점검',
    '8. 파일 구조',
    '9. 운영 및 유지보수 안내',
    '10. 향후 개선 제안',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════
# 1. 프로젝트 개요
# ══════════════════════════════════════════════
doc.add_heading('1. 프로젝트 개요', level=1)

doc.add_heading('1.1 목적', level=2)
doc.add_paragraph(
    '영어조합법인 서풍의 B2B OEM 수산물 가공 사업을 소개하고, '
    '기업의 기술력·품질·비전을 효과적으로 전달하기 위한 기업 홈페이지입니다.'
)

doc.add_heading('1.2 기본 정보', level=2)
add_table(
    ['항목', '내용'],
    [
        ['회사명', '영어조합법인 서풍'],
        ['대표', '상무이사 김태환'],
        ['주소', '전라남도 여수시 석교로 121 (화양면)'],
        ['사업자번호', '417-81-41979'],
        ['호스팅', 'GitHub Pages (정적 사이트)'],
        ['도메인', 'durecoop.github.io/seopung-web'],
    ],
    col_widths=[4, 12],
)

# ══════════════════════════════════════════════
# 2. 기술 사양
# ══════════════════════════════════════════════
doc.add_heading('2. 기술 사양', level=1)

doc.add_heading('2.1 개발 환경', level=2)
add_table(
    ['항목', '사양'],
    [
        ['프레임워크', 'Next.js 16.1.6 (App Router)'],
        ['UI 라이브러리', 'React 19.2.3'],
        ['언어', 'TypeScript 5'],
        ['스타일링', 'Tailwind CSS 4'],
        ['애니메이션', 'GSAP 3.14, Framer Motion 12'],
        ['아이콘', 'Lucide React'],
        ['폰트', 'Pretendard (한국어), Montserrat (영문)'],
        ['배포 방식', '정적 빌드 (Static Export) → GitHub Pages'],
    ],
    col_widths=[4, 12],
)

doc.add_heading('2.2 주요 기술 특징', level=2)
features = [
    ('정적 사이트 생성 (SSG)', '서버 없이 운영 가능, 빠른 로딩 속도'),
    ('반응형 디자인', '모바일 / 태블릿 / 데스크톱 완전 대응'),
    ('스크롤 애니메이션', 'FadeIn 컴포넌트 기반의 부드러운 등장 효과'),
    ('이미지 최적화', 'Next.js Image 컴포넌트 + Sharp 라이브러리 활용'),
    ('SEO 최적화', 'Open Graph, Twitter Card 메타 태그 적용'),
    ('접근성', '시맨틱 HTML, alt 텍스트, 키보드 네비게이션 지원'),
]
for title, desc in features:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title}: ')
    run.bold = True
    p.add_run(desc)

# ══════════════════════════════════════════════
# 3. 사이트 구조
# ══════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('3. 사이트 구조 (전체 12페이지)', level=1)

doc.add_heading('3.1 페이지 구성', level=2)
add_table(
    ['No.', '페이지', '경로', '설명'],
    [
        ['1', '메인 (홈)', '/', '회사 개요, 원료 조달, 공정, 기술, 인증, 제품, 문의 CTA'],
        ['2', '회사소개', '/about', '인사말, 경영수치, 경영철학, 연혁, One Platform, 조직, 회사정보'],
        ['3', '생산공정', '/process', '원료 조달(49호 중매인), 6단계 제조공정 상세'],
        ['4', '기술·설비', '/technology', 'AI X-ray, AI 초분광, 장비, 자동화 로드맵, 투자계획'],
        ['5', '품질·인증', '/certification', 'HACCP, 이력추적, ASC/MSC, 품질인증, FSSC 22000'],
        ['6', '제품', '/products', '4개 카테고리(냉동수산물, 밀키트, 영광굴비, ASC/MSC), 9개 어종'],
        ['7', '영광굴비', '/gulbi', '프리미엄 굴비 소개, 전통 공법, 장인 작업'],
        ['8', '비전', '/vision', '5대 전략, 브랜드 히스토리, 비전 목표'],
        ['9', '문의', '/contact', '문의 폼, 회사정보, 계열사 안내'],
        ['10', '소식', '/news', '최신 뉴스 5건'],
        ['11', '자료실', '/resources', '인증서, 사진 갤러리(18장), 회사 브로슈어'],
        ['12', '관리자', '/admin', '비밀번호 보호, 문의 관리, 공지사항 관리'],
    ],
    col_widths=[1, 3, 3, 9],
)

doc.add_heading('3.2 네비게이션 구조', level=2)
nav = """홈
  ├── 회사소개
  ├── 생산공정
  ├── 기술·설비
  ├── 품질·인증
  ├── 제품
  │     └── 영광굴비
  ├── 비전
  ├── 문의
  ├── 소식
  └── 자료실
  (관리자 - 네비게이션에 미노출)"""
p = doc.add_paragraph()
run = p.add_run(nav)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ══════════════════════════════════════════════
# 3.3 페이지별 스크린샷
# ══════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('3.3 페이지별 스크린샷', level=2)

SS = 'D:/2_Projects/web/Maker_homepade/seopung-web/screenshots'

page_screenshots = [
    ('메인 (홈)', 'home_hero.png', '/'),
    ('회사소개', 'about_hero.png', '/about'),
    ('생산공정', 'process_hero.png', '/process'),
    ('기술·설비', 'technology_hero.png', '/technology'),
    ('품질·인증', 'certification_hero.png', '/certification'),
    ('제품', 'products_hero.png', '/products'),
    ('영광굴비', 'gulbi_hero.png', '/gulbi'),
    ('비전', 'vision_hero.png', '/vision'),
    ('문의', 'contact_hero.png', '/contact'),
    ('소식', 'news_hero.png', '/news'),
    ('자료실', 'resources_hero.png', '/resources'),
]

for label, filename, route in page_screenshots:
    filepath = f'{SS}/{filename}'
    p = doc.add_paragraph()
    run = p.add_run(f'■ {label}')
    run.bold = True
    run.font.size = Pt(11)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f'경로: {route}')
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_picture(filepath, width=Inches(6.2))
    doc.add_paragraph()  # spacing

# Mobile screenshots
doc.add_page_break()
doc.add_heading('3.4 모바일 반응형 스크린샷', level=2)
p = doc.add_paragraph()
run = p.add_run('모바일 기기(390x844px)에서의 화면 표시 상태입니다.')
run.italic = True

mobile_shots = [
    ('메인 (모바일)', 'home_mobile.png'),
    ('회사소개 (모바일)', 'about_mobile.png'),
    ('제품 (모바일)', 'products_mobile.png'),
]

# Place mobile screenshots side by side using a table
mobile_table = doc.add_table(rows=2, cols=3)
mobile_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, filename) in enumerate(mobile_shots):
    # Label row
    cell = mobile_table.rows[0].cells[i]
    cell.text = label
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
    # Image row
    cell = mobile_table.rows[1].cells[i]
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].add_run()
    run.add_picture(f'{SS}/{filename}', width=Inches(1.8))

doc.add_paragraph()

# ══════════════════════════════════════════════
# 4. 주요 기능
# ══════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('4. 주요 기능', level=1)

doc.add_heading('4.1 인터랙티브 기능', level=2)
interactive = [
    '스크롤 감지 네비게이션 바 (배경색/로고 크기 자동 변환)',
    '모바일 햄버거 메뉴',
    '페이지 상단 스크롤 진행률 표시 바 (골드 그라디언트)',
    '스크롤 시 요소 등장 애니메이션',
    '통계 숫자 카운트업 애니메이션',
    '사진 갤러리 라이트박스 (자료실)',
    '카테고리별 탭 필터링',
    '문의 폼 유효성 검사',
    '스크롤 맨 위로 이동 버튼',
]
for item in interactive:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2 관리자 기능', level=2)
admin_items = [
    '비밀번호 기반 접근 제어',
    '문의 접수 관리 (읽음 처리, 삭제)',
    '공지사항 등록/삭제',
    '데이터 저장: 브라우저 localStorage 기반',
]
for item in admin_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.3 디자인 시스템', level=2)
add_table(
    ['요소', '값'],
    [
        ['주요 배경색', 'Navy-950 (짙은 남색)'],
        ['강조색', 'Gold-400/500 (골드)'],
        ['보조색', 'Ocean-400/500 (청록)'],
        ['모서리', '2xl (둥근 카드)'],
        ['카드 스타일', '유리 효과 (backdrop-blur + 반투명 배경)'],
    ],
    col_widths=[4, 12],
)

# ══════════════════════════════════════════════
# 5. 콘텐츠 현황
# ══════════════════════════════════════════════
doc.add_heading('5. 콘텐츠 현황', level=1)

doc.add_heading('5.1 적용된 실제 콘텐츠', level=2)
add_table(
    ['항목', '상태', '세부내용'],
    [
        ['회사 기본정보', '✅ 완료', '상호, 주소, 사업자번호'],
        ['인사말', '✅ 완료', '상무이사 김태환 인사말'],
        ['경영수치', '✅ 완료', '매출 379억(2025), 목표 400억(2026) 등'],
        ['연혁', '✅ 완료', '2008~2026 주요 이력 8건'],
        ['인증정보', '✅ 완료', 'HACCP, 이력추적, ASC/MSC, 품질인증'],
        ['제품정보', '✅ 완료', '4개 카테고리, 9개 어종, 134+ 개발품목'],
        ['파트너사', '✅ 완료', '풀무원, 푸드머스, 홈플러스, 이마트, 쿠팡'],
        ['뉴스', '✅ 완료', '5건 (2023.11 ~ 2026.01)'],
        ['투자계획', '✅ 완료', '2024~2027 연도별 투자내역'],
        ['사진 자료', '✅ 완료', '약 150장 (공장, 공정, 설비, 팀, 제품 등)'],
    ],
    col_widths=[3, 2, 11],
)

doc.add_heading('5.2 임시(더미) 데이터 — 실제 자료로 교체 필요', level=2)
p = doc.add_paragraph()
run = p.add_run('아래 항목들은 현재 임시 데이터가 적용되어 있으며, 실제 자료로 교체가 필요합니다.')
run.italic = True

add_table(
    ['항목', '현재 상태', '위치'],
    [
        ['전화번호', '061-XXX-XXXX (마스킹 처리)', 'Footer, 홈 문의섹션'],
        ['이메일', 'seopung@example.com (임시)', 'Footer'],
        ['SNS 링크', '# (미연결)', 'Footer (LinkedIn, Blog 아이콘)'],
        ['인증서 PDF', '"준비 중" 표시', '자료실 다운로드 버튼'],
        ['회사 브로슈어', '"준비 중" 표시', '자료실 하단'],
        ['관리자 비밀번호', '코드 내 하드코딩', '관리자 페이지'],
    ],
    col_widths=[3, 6, 7],
)

# ══════════════════════════════════════════════
# 6. 요청 자료 목록
# ══════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('6. 요청 자료 목록 (최종 완성에 필요)', level=1)
doc.add_paragraph('아래 자료를 제공해 주시면 홈페이지에 즉시 반영하겠습니다.')

doc.add_heading('6.1 필수 자료 (반드시 필요)', level=2)
add_table(
    ['No.', '항목', '형식', '설명'],
    [
        ['1', '대표 전화번호', '텍스트', '현재 061-XXX-XXXX으로 마스킹된 부분'],
        ['2', '대표 이메일', '텍스트', '현재 seopung@example.com 임시 주소'],
        ['3', '대표 팩스번호', '텍스트', '필요시 추가'],
        ['4', '인사말 사진', 'JPG/PNG (세로 4:5)', '김태환 상무이사 프로필 사진 (교체 희망 시)'],
    ],
    col_widths=[1, 3.5, 4, 7.5],
)

doc.add_heading('6.2 인증서 관련', level=2)
add_table(
    ['No.', '항목', '형식', '설명'],
    [
        ['5', 'HACCP 인증서', 'PDF 또는 이미지', '자료실 다운로드용'],
        ['6', '수산물이력추적 인증서', 'PDF 또는 이미지', '자료실 다운로드용'],
        ['7', 'ASC/MSC 인증서', 'PDF 또는 이미지', '자료실 다운로드용'],
        ['8', '품질인증서', 'PDF 또는 이미지', '자료실 다운로드용'],
        ['9', '사업자등록증', 'PDF 또는 이미지', '필요시 자료실 추가'],
    ],
    col_widths=[1, 4.5, 4, 6.5],
)

doc.add_heading('6.3 홍보물', level=2)
add_table(
    ['No.', '항목', '형식', '설명'],
    [
        ['10', '회사 브로슈어', 'PDF', '자료실 다운로드 파일로 등록'],
        ['11', '제품 카탈로그', 'PDF', '제품 페이지 또는 자료실에 추가 가능'],
    ],
    col_widths=[1, 3.5, 4.5, 7],
)

doc.add_heading('6.4 제품 패키지 사진 (핵심 요청)', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '현재 제품 페이지에 실제 패키지 사진이 2장(풀무원 ASC 새우살, 밀키트)만 있습니다. '
    '서풍의 주력 상품과 자랑할 만한 제품을 중심으로 아래 사진을 요청드립니다.'
)
run.italic = True

p = doc.add_paragraph()
run = p.add_run('촬영 권장사항: ')
run.bold = True
p.add_run('흰색 또는 단색 배경, 정면 촬영, 최소 1200x1200px 이상, JPG 또는 PNG')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('■ 냉동수산가공 카테고리 (현재: 공장 설비 사진으로 대체 중)')
run.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

add_table(
    ['No.', '제품명', '현재 상태', '요청 사항'],
    [
        ['12', '고등어필렛 패키지', '사진 없음 (터널프리저 설비 사진으로 대체)', '실제 제품 패키지 정면 사진'],
        ['13', '삼치필렛 패키지', '사진 없음', '실제 제품 패키지 정면 사진'],
        ['14', '갈치손질 패키지', '사진 없음', '실제 제품 패키지 정면 사진'],
        ['15', '오징어손질 패키지', '사진 없음', '실제 제품 패키지 정면 사진'],
        ['16', '아귀손질 패키지', '사진 없음', '실제 제품 패키지 정면 사진'],
    ],
    col_widths=[1, 3.5, 5.5, 6],
)

p = doc.add_paragraph()
run = p.add_run('■ 밀키트·간편식 카테고리 (현재: 밀키트 1장만 있음)')
run.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

add_table(
    ['No.', '제품명', '현재 상태', '요청 사항'],
    [
        ['17', '수산물밀키트 추가', '1장 있음 (박스+트레이 사진)', '다양한 밀키트 종류별 패키지 사진 2~3장'],
        ['18', '간편조리수산물', '사진 없음', '실제 제품 패키지 사진'],
        ['19', '양념수산물', '사진 없음', '실제 제품 패키지 사진'],
    ],
    col_widths=[1, 3.5, 5.5, 6],
)

p = doc.add_paragraph()
run = p.add_run('■ 프리미엄 영광굴비 카테고리 (현재: 건조 작업 사진으로 대체)')
run.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

add_table(
    ['No.', '제품명', '현재 상태', '요청 사항'],
    [
        ['20', '참조기굴비세트', '사진 없음 (건조장 작업 사진 대체)', '완성된 선물세트 패키지 사진'],
        ['21', '보리굴비', '사진 없음', '완성 제품 사진 또는 패키지'],
        ['22', '선물세트 박스', '사진 없음', '고급 선물 포장 박스 정면 사진'],
    ],
    col_widths=[1, 3.5, 5.5, 6],
)

p = doc.add_paragraph()
run = p.add_run('■ ASC/MSC 인증 제품 (현재: 풀무원 ASC 새우살 1장)')
run.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

add_table(
    ['No.', '제품명', '현재 상태', '요청 사항'],
    [
        ['23', '서풍 자체 ASC 인증 제품', '풀무원 PB 사진만 있음', '서풍 브랜드 또는 OEM 납품 제품 패키지 사진'],
        ['24', 'MSC 인증 제품', '사진 없음', 'MSC 마크가 보이는 제품 패키지 사진'],
    ],
    col_widths=[1, 4, 5, 6],
)

doc.add_heading('6.5 자랑할 만한 콘텐츠 추가 요청', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '서풍의 강점을 더 효과적으로 보여줄 수 있는 자료가 있으시면 제공 부탁드립니다.'
)

add_table(
    ['No.', '항목', '설명'],
    [
        ['25', '주력 상품 Best 3~5', '서풍에서 가장 자신 있는 주력 제품이 무엇인지 알려주시면 메인/제품 페이지에 별도 하이라이트 섹션으로 강조 가능'],
        ['26', '수상 실적/표창장', '우수업체 선정, 품질인증 표창 등 수상 사진 또는 스캔본'],
        ['27', '납품 실적 사진', '대형 유통사 매대에 진열된 서풍 OEM 제품 사진 (풀무원, 이마트 PB 등)'],
        ['28', '공장 전경/항공 사진', '공장 외관 또는 드론 촬영 사진 (현재 내부 사진만 있음)'],
        ['29', '대표/임원 프로필 사진', '김태환 상무이사 공식 프로필 사진 (현재 공장 작업 사진 대체)'],
        ['30', '파트너사 협약/MOU 사진', '주요 파트너사와의 계약 체결, 방문 사진 등'],
    ],
    col_widths=[1, 4, 11],
)

doc.add_heading('6.6 선택 자료', level=2)
add_table(
    ['No.', '항목', '형식', '설명'],
    [
        ['31', 'SNS 계정 주소', 'URL', '블로그, 인스타그램 등 운영 중인 채널'],
        ['32', '파트너사 로고', 'PNG (투명배경)', '풀무원, 홈플러스 등 (현재 텍스트만 표시)'],
        ['33', '지도 좌표', '위도/경도', '문의 페이지에 지도 삽입 희망 시'],
        ['34', '독자 도메인', 'URL', '예: www.seopung.co.kr'],
        ['35', '관리자 비밀번호', '텍스트', '운영용 비밀번호 지정'],
    ],
    col_widths=[1, 3.5, 4, 7.5],
)

# ══════════════════════════════════════════════
# 7. 기존 사진 퀄리티 점검
# ══════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('7. 기존 사진 퀄리티 점검', level=1)
doc.add_paragraph(
    '현재 홈페이지에 사용 중인 약 100장의 사진을 점검한 결과, '
    '아래 항목들의 교체 또는 보완을 권장합니다.'
)

doc.add_heading('7.1 해상도 부족 (교체 권장)', level=2)
p = doc.add_paragraph()
run = p.add_run('아래 이미지들은 파일 크기가 80~150KB 수준으로, 확대 시 화질이 떨어집니다.')
run.italic = True

add_table(
    ['파일명', '크기', '사용 위치', '문제점'],
    [
        ['ai-xray-process.png', '105KB', '기술·설비 페이지 (AI X-ray 공정도)', '해상도 매우 낮음, 텍스트가 흐릿하게 보임. 고해상도 원본 또는 재제작 필요'],
        ['ai-xray-diagram.png', '152KB', '기술·설비 페이지 (AI X-ray 시스템 구성도)', '저해상도, 글자가 작아서 읽기 어려움. 원본 제공 또는 재제작 필요'],
        ['radiation-tester.jpg', '112KB', '기술·설비 페이지 (방사능 검사기)', '해상도 낮음, 장비 디테일이 잘 안 보임'],
        ['fishtrace-logo-2017.jpg', '89KB', '인증 페이지 (이력추적 스티커 레이아웃)', '인쇄용 스티커 레이아웃 이미지, 홈페이지용으로 부적합'],
    ],
    col_widths=[3.5, 1.5, 4, 7],
)

doc.add_heading('7.2 용도 부적합 (보완 권장)', level=2)
add_table(
    ['파일명', '사용 위치', '문제점 및 제안'],
    [
        ['04-tunnel-freezer.jpg', '제품 페이지 히어로, 홈 공정 섹션 등 5곳', '동일 사진이 5개 이상 페이지에 반복 사용됨. 제품 페이지에는 실제 완제품 사진이 더 적합'],
        ['06-cold-storage.jpg', '제품 CTA 섹션, 홈 제품 티저', '냉동창고 내부 사진이 제품 소개 영역에 사용 중. 제품 라인업 사진으로 교체 권장'],
        ['drying-rack.jpg', '제품 > 영광굴비 카테고리 대표 이미지', '작업 중 사진(주황 통에서 세척). 완성된 굴비 건조대 전경 사진이 더 효과적'],
        ['safety-sign.jpg', '자료실 갤러리', '공장 안전 표지판 사진. 퀄리티 관점에서 대외 공개용으로 부적합'],
    ],
    col_widths=[3.5, 4, 8.5],
)

doc.add_heading('7.3 교체하면 더 좋은 사진', level=2)
add_table(
    ['현재 사진', '위치', '교체 제안'],
    [
        ['director-writing.jpg (공장에서 메모하는 사진)', '회사소개 인사말 섹션', '김태환 상무이사 공식 프로필 사진 또는 정장 착용 상반신 사진'],
        ['factory-team.jpg (현장 단체사진)', '회사소개 히어로, 비전 페이지', '전직원 단체사진 (공장 앞 또는 회의실, 밝은 조명)'],
        ['공장 내부 사진 전반', '생산공정 페이지', '일부 사진이 어두움. 조명이 밝은 환경에서 재촬영하면 청결함이 더 부각됨'],
    ],
    col_widths=[4, 3.5, 8.5],
)

doc.add_heading('7.4 사진 현황 요약', level=2)
add_table(
    ['카테고리', '보유 수량', '품질 평가', '비고'],
    [
        ['공정(원료~보관)', '33장', '양호', '6단계 공정 모두 커버됨. 일부 어두운 사진 교체 권장'],
        ['경매/조달', '16장', '양호', '49호 중매인, 경매장, 원료 검수 등 충분'],
        ['설비/기술', '14장', '보통', 'AI 공정도 2장 저해상도 → 교체 필요'],
        ['영광굴비', '11장', '양호', '전통 공법 단계별 잘 촬영됨'],
        ['팀/인물', '14장', '양호', '다양한 장면 확보, 대표 프로필 사진만 부족'],
        ['인증/로고', '7장', '보통', '스티커 레이아웃 이미지 부적합, HACCP/ASC/MSC 로고 PNG 필요'],
        ['제품 패키지', '2장', '부족', '풀무원 ASC 1장, 밀키트 1장뿐. 최소 10장 이상 필요'],
        ['히어로/배경', '11장', '양호', '새벽 항구, 경매장 등 분위기 좋음'],
    ],
    col_widths=[3, 2, 2, 9],
)

# ══════════════════════════════════════════════
# 8. 파일 구조
# ══════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('8. 파일 구조', level=1)

structure = """seopung-web/
├── src/
│   ├── app/                        페이지 (12개)
│   │   ├── page.tsx                메인(홈)
│   │   ├── layout.tsx              루트 레이아웃
│   │   ├── about/page.tsx          회사소개
│   │   ├── process/page.tsx        생산공정
│   │   ├── technology/page.tsx     기술·설비
│   │   ├── certification/page.tsx  품질·인증
│   │   ├── products/page.tsx       제품
│   │   ├── gulbi/page.tsx          영광굴비
│   │   ├── vision/page.tsx         비전
│   │   ├── contact/page.tsx        문의
│   │   ├── news/page.tsx           소식
│   │   ├── resources/page.tsx      자료실
│   │   └── admin/page.tsx          관리자
│   ├── components/                 재사용 컴포넌트
│   │   ├── ui/                     Navbar, Footer, Breadcrumb 등
│   │   └── sections/               홈 섹션 컴포넌트
│   ├── hooks/                      커스텀 훅
│   └── lib/                        유틸리티, 관리자 스토어
├── public/images/                  이미지 약 150장
├── package.json                    의존성 관리
├── next.config.ts                  Next.js 설정
└── tsconfig.json                   TypeScript 설정"""

p = doc.add_paragraph()
run = p.add_run(structure)
run.font.name = 'Consolas'
run.font.size = Pt(8.5)

# ══════════════════════════════════════════════
# 9. 운영 및 유지보수
# ══════════════════════════════════════════════
doc.add_heading('9. 운영 및 유지보수 안내', level=1)

doc.add_heading('9.1 콘텐츠 수정', level=2)
doc.add_paragraph('각 페이지의 텍스트는 src/app/[페이지명]/page.tsx 파일에서 직접 수정합니다.', style='List Bullet')
doc.add_paragraph('이미지 교체: public/images/ 폴더에 동일 파일명으로 교체합니다.', style='List Bullet')

doc.add_heading('9.2 빌드 및 배포', level=2)
doc.add_paragraph('main 브랜치에 푸시하면 GitHub Actions가 자동으로 빌드·배포합니다.')
p = doc.add_paragraph()
run = p.add_run('배포 절차:\n1. 소스 수정\n2. npm run build (로컬 빌드 확인)\n3. git add → git commit → git push origin main\n4. 약 1~2분 후 자동 배포 완료')
run.font.size = Pt(9)

doc.add_heading('9.3 관리자 페이지', level=2)
doc.add_paragraph('접속: https://durecoop.github.io/seopung-web/admin', style='List Bullet')
doc.add_paragraph('기능: 문의 접수 확인, 공지사항 관리', style='List Bullet')
doc.add_paragraph('주의: localStorage 기반이므로 브라우저/기기별 데이터 독립', style='List Bullet')

doc.add_heading('9.4 주의사항', level=2)
warnings = [
    '정적 사이트이므로 문의 폼 데이터는 서버에 저장되지 않고 해당 브라우저의 localStorage에만 저장됩니다.',
    '실제 문의 접수를 위해서는 별도의 백엔드 서버 또는 이메일 연동이 필요합니다.',
    '관리자 비밀번호가 소스 코드에 포함되어 있으므로, 운영 전 변경을 권장합니다.',
]
for w in warnings:
    p = doc.add_paragraph()
    run = p.add_run(f'⚠ {w}')
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run.font.size = Pt(9)

# ══════════════════════════════════════════════
# 10. 향후 개선 제안
# ══════════════════════════════════════════════
doc.add_heading('10. 향후 개선 제안', level=1)
add_table(
    ['우선순위', '항목', '설명'],
    [
        ['높음', '문의 폼 백엔드 연동', '이메일 발송 또는 서버 저장 (현재 localStorage만 사용)'],
        ['높음', '독자 도메인 연결', '자체 도메인(예: seopung.co.kr) 연결'],
        ['중간', '제품 상세 페이지', '개별 제품 클릭 시 상세 정보 페이지'],
        ['중간', '다국어 지원', '영문/일문 페이지 (해외 바이어 대응)'],
        ['낮음', '찾아오시는 길 지도', '카카오맵/네이버지도 연동'],
        ['낮음', '뉴스 관리 기능', '관리자에서 뉴스 추가/수정 기능'],
    ],
    col_widths=[2.5, 4.5, 9],
)

# ── 서명란 ──
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 끝 —')
run.font.size = Pt(12)
run.bold = True

doc.add_paragraph()

sign_table = doc.add_table(rows=2, cols=2)
sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sign_table.style = 'Table Grid'
sign_table.rows[0].cells[0].text = '납품자'
sign_table.rows[0].cells[1].text = '인수자'
sign_table.rows[1].cells[0].text = '\n\n두레쿱\n\n서명: ________________'
sign_table.rows[1].cells[1].text = '\n\n영어조합법인 서풍\n\n서명: ________________'

for row in sign_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 저장 ──
output = 'D:/2_Projects/web/Maker_homepade/seopung-web/납품설명서_서풍홈페이지.docx'
doc.save(output)
print(f'문서 생성 완료: {output}')
